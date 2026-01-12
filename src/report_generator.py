"""
Módulo para generación de informes en PDF y DOCX
"""
import os
from typing import Dict, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors


class ReportGenerator:
    """Genera informes de evaluación en diferentes formatos"""
    
    def __init__(self, output_format: str = "pdf"):
        """
        Inicializa el generador de informes
        
        Args:
            output_format: Formato de salida ('pdf' o 'docx')
        """
        self.output_format = output_format.lower()
    
    def generate_author_report(
        self,
        file_path: str,
        evaluation: Dict[str, List[str]],
        manuscript_title: Optional[str] = None
    ) -> str:
        """
        Genera informe para el autor (solo evaluación)
        
        Args:
            file_path: Ruta base para el archivo de salida
            evaluation: Diccionario con evaluación
            manuscript_title: Título del manuscrito (opcional)
            
        Returns:
            Ruta del archivo generado
        """
        base_name = os.path.splitext(file_path)[0]
        output_path = f"{base_name}_Author_Report.{self.output_format}"
        
        if self.output_format == 'pdf':
            self._generate_author_pdf(output_path, evaluation, manuscript_title)
        else:  # docx
            self._generate_author_docx(output_path, evaluation, manuscript_title)
        
        return output_path
    
    def generate_auditor_report(
        self,
        file_path: str,
        evaluation: Dict[str, List[str]],
        pubmed_data: Dict[str, List[Dict]],
        keyphrases: List[str],
        manuscript_text: str,
        article_type: str,
        manuscript_title: Optional[str] = None
    ) -> str:
        """
        Genera informe completo para auditoría
        
        Args:
            file_path: Ruta base para el archivo de salida
            evaluation: Diccionario con evaluación
            pubmed_data: Datos de PubMed
            keyphrases: Frases clave extraídas
            manuscript_text: Texto del manuscrito
            article_type: Tipo de artículo
            manuscript_title: Título del manuscrito (opcional)
            
        Returns:
            Ruta del archivo generado
        """
        base_name = os.path.splitext(file_path)[0]
        output_path = f"{base_name}_Auditor_Report.{self.output_format}"
        
        if self.output_format == 'pdf':
            self._generate_auditor_pdf(
                output_path, evaluation, pubmed_data, keyphrases,
                manuscript_text, article_type, manuscript_title
            )
        else:  # docx
            self._generate_auditor_docx(
                output_path, evaluation, pubmed_data, keyphrases,
                manuscript_text, article_type, manuscript_title
            )
        
        return output_path
    
    def _generate_author_pdf(
        self,
        output_path: str,
        evaluation: Dict[str, List[str]],
        manuscript_title: Optional[str]
    ):
        """Genera informe PDF para el autor"""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Estilos personalizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a5490'),
            spaceAfter=30,
            alignment=1  # Centrado
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2e75b6'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Título
        story.append(Paragraph("Peer Review Report", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        if manuscript_title:
            story.append(Paragraph(f"<b>Manuscript:</b> {manuscript_title}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%Y-%m-%d')}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Secciones de evaluación
        section_titles = {
            'major': 'Major Points',
            'minor': 'Minor Points',
            'other': 'Other Points',
            'suggestions': 'Suggestions for Improvement'
        }
        
        for key, title in section_titles.items():
            if evaluation.get(key):
                story.append(Paragraph(title, heading_style))
                for point in evaluation[key]:
                    story.append(Paragraph(f"• {point}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
    
    def _generate_author_docx(
        self,
        output_path: str,
        evaluation: Dict[str, List[str]],
        manuscript_title: Optional[str]
    ):
        """Genera informe DOCX para el autor"""
        doc = Document()
        
        # Título
        title = doc.add_heading('Peer Review Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información básica
        if manuscript_title:
            p = doc.add_paragraph()
            p.add_run('Manuscript: ').bold = True
            p.add_run(manuscript_title)
        
        p = doc.add_paragraph()
        p.add_run('Date: ').bold = True
        p.add_run(datetime.now().strftime('%Y-%m-%d'))
        
        doc.add_paragraph()  # Espacio
        
        # Secciones de evaluación
        section_titles = {
            'major': 'Major Points',
            'minor': 'Minor Points',
            'other': 'Other Points',
            'suggestions': 'Suggestions for Improvement'
        }
        
        for key, title in section_titles.items():
            if evaluation.get(key):
                doc.add_heading(title, level=1)
                for point in evaluation[key]:
                    doc.add_paragraph(point, style='List Bullet')
        
        doc.save(output_path)
    
    def _generate_auditor_pdf(
        self,
        output_path: str,
        evaluation: Dict[str, List[str]],
        pubmed_data: Dict[str, List[Dict]],
        keyphrases: List[str],
        manuscript_text: str,
        article_type: str,
        manuscript_title: Optional[str]
    ):
        """Genera informe PDF completo para auditoría"""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Estilos
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a5490'),
            spaceAfter=30,
            alignment=1
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2e75b6'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Título
        story.append(Paragraph("Auditor Report - Detailed Analysis", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Información del manuscrito
        story.append(Paragraph("Manuscript Information", heading_style))
        if manuscript_title:
            story.append(Paragraph(f"<b>Title:</b> {manuscript_title}", styles['Normal']))
        story.append(Paragraph(f"<b>Type:</b> {article_type}", styles['Normal']))
        story.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%Y-%m-%d')}", styles['Normal']))
        story.append(Paragraph(f"<b>Length:</b> {len(manuscript_text)} characters", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Frases clave
        story.append(Paragraph("Extracted Key Phrases", heading_style))
        for kp in keyphrases:
            story.append(Paragraph(f"• {kp}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Resultados de PubMed
        story.append(Paragraph("PubMed Search Results", heading_style))
        total_articles = sum(len(articles) for articles in pubmed_data.values())
        story.append(Paragraph(f"<b>Total articles retrieved:</b> {total_articles}", styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        for keyphrase, articles in pubmed_data.items():
            story.append(Paragraph(f"<b>{keyphrase}:</b> {len(articles)} articles", styles['Normal']))
            for i, art in enumerate(articles[:3], 1):  # Solo primeros 3
                story.append(Paragraph(
                    f"  {i}. {art.get('title', 'No title')} ({art.get('year', 'N/A')})",
                    styles['Normal']
                ))
        story.append(Spacer(1, 0.3*inch))
        
        # Evaluación
        story.append(PageBreak())
        story.append(Paragraph("Evaluation Results", heading_style))
        
        section_titles = {
            'major': 'Major Points',
            'minor': 'Minor Points',
            'other': 'Other Points',
            'suggestions': 'Suggestions for Improvement'
        }
        
        for key, title in section_titles.items():
            if evaluation.get(key):
                story.append(Paragraph(title, heading_style))
                for point in evaluation[key]:
                    story.append(Paragraph(f"• {point}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
    
    def _generate_auditor_docx(
        self,
        output_path: str,
        evaluation: Dict[str, List[str]],
        pubmed_data: Dict[str, List[Dict]],
        keyphrases: List[str],
        manuscript_text: str,
        article_type: str,
        manuscript_title: Optional[str]
    ):
        """Genera informe DOCX completo para auditoría"""
        doc = Document()
        
        # Título
        title = doc.add_heading('Auditor Report - Detailed Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del manuscrito
        doc.add_heading('Manuscript Information', level=1)
        if manuscript_title:
            p = doc.add_paragraph()
            p.add_run('Title: ').bold = True
            p.add_run(manuscript_title)
        
        p = doc.add_paragraph()
        p.add_run('Type: ').bold = True
        p.add_run(article_type)
        
        p = doc.add_paragraph()
        p.add_run('Date: ').bold = True
        p.add_run(datetime.now().strftime('%Y-%m-%d'))
        
        p = doc.add_paragraph()
        p.add_run('Length: ').bold = True
        p.add_run(f'{len(manuscript_text)} characters')
        
        # Frases clave
        doc.add_heading('Extracted Key Phrases', level=1)
        for kp in keyphrases:
            doc.add_paragraph(kp, style='List Bullet')
        
        # Resultados de PubMed
        doc.add_heading('PubMed Search Results', level=1)
        total_articles = sum(len(articles) for articles in pubmed_data.values())
        p = doc.add_paragraph()
        p.add_run('Total articles retrieved: ').bold = True
        p.add_run(str(total_articles))
        
        for keyphrase, articles in pubmed_data.items():
            doc.add_heading(f'{keyphrase}: {len(articles)} articles', level=2)
            for i, art in enumerate(articles[:5], 1):  # Primeros 5
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{art.get('title', 'No title')} ").bold = True
                p.add_run(f"({art.get('year', 'N/A')})\n")
                p.add_run(f"Authors: {art.get('authors', 'Unknown')}\n")
                p.add_run(f"Journal: {art.get('journal', 'Unknown')}")
        
        # Evaluación
        doc.add_page_break()
        doc.add_heading('Evaluation Results', level=1)
        
        section_titles = {
            'major': 'Major Points',
            'minor': 'Minor Points',
            'other': 'Other Points',
            'suggestions': 'Suggestions for Improvement'
        }
        
        for key, title in section_titles.items():
            if evaluation.get(key):
                doc.add_heading(title, level=2)
                for point in evaluation[key]:
                    doc.add_paragraph(point, style='List Bullet')
        
        doc.save(output_path)
