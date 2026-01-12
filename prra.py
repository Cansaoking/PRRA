import sys
import json
import os
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit,
                             QPushButton, QTextEdit, QComboBox, QCheckBox, QFileDialog, QProgressBar, QTabWidget,
                             QMessageBox, QPlainTextEdit)
from PyQt5.QtCore import QThread, pyqtSignal
import torch
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM
from docx import Document
from PyPDF2 import PdfReader
from striprtf.striprtf import rtf_to_text
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from Bio import Entrez
import datetime
import re

Entrez.email = "tuemail@example.com"  # Requerido por Entrez para rastreo (usa uno real)

class IAWorker(QThread):
    progress = pyqtSignal(int)
    result = pyqtSignal(dict)
    error = pyqtSignal(str)

    def __init__(self, file_path, num_keys, num_articles, model_name, prompts, manual_mode, output_format):
        super().__init__()
        self.file_path = file_path
        self.num_keys = num_keys
        self.num_articles = num_articles
        self.model_name = model_name
        self.prompts = prompts
        self.manual_mode = manual_mode
        self.output_format = output_format
        self.device = "cuda" if torch.cuda.is_available() else "cpu"

    def run(self):
        try:
            # Paso 1: Extraer texto del manuscrito
            text = self.extract_text(self.file_path)
            self.progress.emit(10)

            # Paso 2: Detectar tipo de artículo (simple heurística)
            article_type = self.detect_article_type(text)

            # Paso 3: Extraer frases clave con IA
            model, tokenizer = self.load_model(self.model_name)
            keyphrases = self.extract_keyphrases(text, tokenizer, model, self.num_keys)
            if self.manual_mode:
                # Aquí simular confirmación; en UI real, emitir señal para diálogo
                pass
            self.progress.emit(30)

            # Paso 4: Buscar en PubMed
            pubmed_data = self.search_pubmed(keyphrases, self.num_articles)
            if self.manual_mode:
                pass
            self.progress.emit(60)

            # Paso 5: Analizar con IA usando abstracts
            evaluation = self.analyze_manuscript(text, pubmed_data, tokenizer, model, article_type)
            self.progress.emit(80)

            # Paso 6: Generar informes
            self.generate_reports(evaluation, pubmed_data, keyphrases, text, self.output_format)
            self.progress.emit(100)

            self.result.emit({"success": True})
        except Exception as e:
            self.error.emit(str(e))

    def extract_text(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            reader = PdfReader(file_path)
            return ' '.join(page.extract_text() for page in reader.pages if page.extract_text())
        elif ext in ['.doc', '.docx']:
            doc = Document(file_path)
            return ' '.join(p.text for p in doc.paragraphs)
        elif ext == '.rtf':
            with open(file_path, 'r') as f:
                return rtf_to_text(f.read())
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        raise ValueError("Formato no soportado")

    def detect_article_type(self, text):
        # Heurística simple: Buscar secciones típicas
        if all(re.search(r'\b(introduction|methods|results|discussion)\b', text, re.I) for _ in range(1)):
            return "Investigación"
        return "Review u otro"

    def load_model(self, model_name):
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForCausalLM.from_pretrained(model_name).to(self.device)
        return model, tokenizer

    def extract_keyphrases(self, text, tokenizer, model, num_keys):
        prompt = self.prompts['keyphrases'].format(num=num_keys, text=text[:2000])  # Limitar para eficiencia
        inputs = tokenizer(prompt, return_tensors="pt").to(self.device)
        outputs = model.generate(**inputs, max_new_tokens=200)
        keyphrases = tokenizer.decode(outputs[0], skip_special_tokens=True).split('\n')[:num_keys]
        return [kp.strip() for kp in keyphrases if kp.strip()]

    def search_pubmed(self, keyphrases, num_articles):
        pubmed_data = {}
        current_year = datetime.datetime.now().year
        for kp in keyphrases:
            query = kp  # Inicial individual
            handle = Entrez.esearch(db="pubmed", term=query, sort="pub date", retmax=100, datetype="pdat", mindate=f"{current_year-5}/01/01", maxdate=f"{current_year}/12/31")
            record = Entrez.read(handle)
            ids = record['IdList']
            if len(ids) > 100:  # Demasiados: combinar con AND si hay más frases
                query = ' AND '.join(keyphrases[:2])  # Ejemplo simple
                handle = Entrez.esearch(db="pubmed", term=query, sort="pub date", retmax=num_articles)
                record = Entrez.read(handle)
                ids = record['IdList']
            elif len(ids) < num_articles:  # Pocos: extender años
                handle = Entrez.esearch(db="pubmed", term=query, sort="pub date", retmax=num_articles)
                record = Entrez.read(handle)
                ids = record['IdList']
            if not ids:
                continue  # O loguear alternativa

            fetch_handle = Entrez.efetch(db="pubmed", id=ids, retmode="xml")
            articles = Entrez.read(fetch_handle)
            pubmed_data[kp] = []
            for art in articles['PubmedArticle'][:num_articles]:
                info = {
                    'title': art['MedlineCitation']['Article']['ArticleTitle'],
                    'authors': ', '.join(a['ForeName'] + ' ' + a['LastName'] for a in art['MedlineCitation']['Article']['AuthorList']),
                    'journal': art['MedlineCitation']['Article']['Journal']['Title'],
                    'year': art['MedlineCitation']['Article']['Journal']['JournalIssue']['PubDate'].get('Year', 'N/A'),
                    'abstract': art['MedlineCitation']['Article'].get('Abstract', {}).get('AbstractText', ['N/A'])[0]
                }
                pubmed_data[kp].append(info)
        return pubmed_data

    def analyze_manuscript(self, text, pubmed_data, tokenizer, model, article_type):
        abstracts = '\n'.join([f"Abstract: {art['abstract']}" for kp_data in pubmed_data.values() for art in kp_data])
        prompt = self.prompts['analysis'].format(text=text[:5000], abstracts=abstracts, type=article_type)
        inputs = tokenizer(prompt, return_tensors="pt").to(self.device)
        outputs = model.generate(**inputs, max_new_tokens=1000)
        eval_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
        # Parsear a estructura: Major, Minor, Other, Sugerencias
        sections = {'major': [], 'minor': [], 'other': [], 'suggestions': []}
        # Lógica simple de parsing (puedes mejorar con regex)
        current_section = None
        for line in eval_text.split('\n'):
            if 'Major' in line: current_section = 'major'
            elif 'Minor' in line: current_section = 'minor'
            elif 'Other' in line: current_section = 'other'
            elif 'Suggestions' in line: current_section = 'suggestions'
            elif current_section:
                sections[current_section].append(line.strip())
        return sections

    def generate_reports(self, evaluation, pubmed_data, keyphrases, text, format):
        base_name = os.path.splitext(self.file_path)[0]
        # Informe para Autor: Solo evaluación
        self.generate_report(base_name + '_Autor.' + format, evaluation, include_details=False, format=format)
        # Informe Auditoría: Todo
        self.generate_report(base_name + '_Auditoria.' + format, evaluation, include_details=True, pubmed_data=pubmed_data, keyphrases=keyphrases, text=text, format=format)

    def generate_report(self, file_path, evaluation, include_details=False, pubmed_data=None, keyphrases=None, text=None, format='pdf'):
        if format == 'pdf':
            c = canvas.Canvas(file_path, pagesize=letter)
            y = 750
            c.drawString(100, y, "Evaluación del Manuscrito")
            y -= 20
            for section, points in evaluation.items():
                c.drawString(100, y, f"{section.capitalize()} Points:")
                y -= 20
                for p in points:
                    c.drawString(120, y, p)
                    y -= 15
            if include_details:
                y -= 20
                c.drawString(100, y, "Detalles Adicionales:")
                y -= 20
                c.drawString(120, y, f"Frases Clave: {', '.join(keyphrases)}")
                y -= 15
                for kp, arts in pubmed_data.items():
                    c.drawString(120, y, f"Para '{kp}': {len(arts)} artículos encontrados.")
                    y -= 15
                # Añadir más si cabe
            c.save()
        elif format == 'docx':
            doc = Document()
            doc.add_heading('Evaluación del Manuscrito', 0)
            for section, points in evaluation.items():
                doc.add_heading(section.capitalize() + ' Points', level=1)
                for p in points:
                    doc.add_paragraph(p)
            if include_details:
                doc.add_heading('Detalles Adicionales', level=1)
                doc.add_paragraph(f"Frases Clave: {', '.join(keyphrases)}")
                for kp, arts in pubmed_data.items():
                    doc.add_paragraph(f"Para '{kp}': {len(arts)} artículos encontrados.")
            doc.save(file_path)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Revisión por Pares Automática")
        self.setGeometry(100, 100, 800, 600)

        central = QWidget()
        layout = QVBoxLayout()

        tabs = QTabWidget()
        # Pestaña Entrada
        input_tab = QWidget()
        input_layout = QVBoxLayout()
        self.file_label = QLabel("Archivo: Ninguno seleccionado")
        btn_open = QPushButton("Abrir Manuscrito")
        btn_open.clicked.connect(self.open_file)
        input_layout.addWidget(self.file_label)
        input_layout.addWidget(btn_open)
        self.preview = QTextEdit()
        self.preview.setReadOnly(True)
        input_layout.addWidget(QLabel("Vista Previa:"))
        input_layout.addWidget(self.preview)
        input_tab.setLayout(input_layout)
        tabs.addTab(input_tab, "Entrada")

        # Pestaña Configuración
        config_tab = QWidget()
        config_layout = QVBoxLayout()
        self.num_keys_edit = QLineEdit("5")
        config_layout.addWidget(QLabel("Número de Frases Clave:"))
        config_layout.addWidget(self.num_keys_edit)
        self.num_articles_edit = QLineEdit("20")
        config_layout.addWidget(QLabel("Número de Artículos por Búsqueda:"))
        config_layout.addWidget(self.num_articles_edit)
        self.model_combo = QComboBox()
        self.model_combo.addItems(["Qwen/Qwen1.5-7B-Chat", "deepseek-ai/deepseek-coder-7b-instruct"])
        config_layout.addWidget(QLabel("Modelo IA:"))
        config_layout.addWidget(self.model_combo)
        self.manual_checkbox = QCheckBox("Modo Manual (Confirmar Pasos)")
        config_layout.addWidget(self.manual_checkbox)
        self.output_combo = QComboBox()
        self.output_combo.addItems(["pdf", "docx"])
        config_layout.addWidget(QLabel("Formato de Salida:"))
        config_layout.addWidget(self.output_combo)

        # Editor de Prompts
        self.prompt_editor = QPlainTextEdit()
        self.default_prompts = {
            'keyphrases': "Extrae {num} frases clave relevantes del siguiente texto: {text}",
            'analysis': "Analiza el manuscrito: {text}. Usa estos abstracts como referencia: {abstracts}. Tipo: {type}. Evalúa: calidad inglés, orden, actualización, originalidad, consistencia, metodología. Estructura: Major Points, Minor Points, Other Points, Suggestions."
        }
        self.prompt_editor.setPlainText(json.dumps(self.default_prompts, indent=4))
        config_layout.addWidget(QLabel("Editor de Prompts (JSON):"))
        config_layout.addWidget(self.prompt_editor)
        btn_save_prompts = QPushButton("Guardar Prompts")
        btn_save_prompts.clicked.connect(self.save_prompts)
        btn_load_prompts = QPushButton("Cargar Prompts")
        btn_load_prompts.clicked.connect(self.load_prompts)
        prompt_btn_layout = QHBoxLayout()
        prompt_btn_layout.addWidget(btn_save_prompts)
        prompt_btn_layout.addWidget(btn_load_prompts)
        config_layout.addLayout(prompt_btn_layout)
        config_tab.setLayout(config_layout)
        tabs.addTab(config_tab, "Configuración")

        # Pestaña Progreso
        progress_tab = QWidget()
        progress_layout = QVBoxLayout()
        self.progress_bar = QProgressBar()
        self.progress_bar.setMaximum(100)
        progress_layout.addWidget(self.progress_bar)
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        progress_layout.addWidget(self.log_text)
        progress_tab.setLayout(progress_layout)
        tabs.addTab(progress_tab, "Progreso")

        layout.addWidget(tabs)
        btn_start = QPushButton("Iniciar Revisión")
        btn_start.clicked.connect(self.start_review)
        layout.addWidget(btn_start)

        central.setLayout(layout)
        self.setCentralWidget(central)

        self.file_path = None
        self.prompts = self.default_prompts

    def open_file(self):
        self.file_path, _ = QFileDialog.getOpenFileName(self, "Abrir Manuscrito", "", "Documentos (*.doc *.docx *.pdf *.rtf *.txt)")
        if self.file_path:
            self.file_label.setText(f"Archivo: {os.path.basename(self.file_path)}")
            # Vista previa (primeros 1000 chars)
            try:
                text = IAWorker.extract_text(None, self.file_path)  # Llamada estática
                self.preview.setText(text[:1000] + "...")
            except:
                QMessageBox.warning(self, "Error", "No se pudo leer el archivo.")

    def save_prompts(self):
        file, _ = QFileDialog.getSaveFileName(self, "Guardar Prompts", "", "JSON (*.json)")
        if file:
            with open(file, 'w') as f:
                json.dump(self.prompts, f)

    def load_prompts(self):
        file, _ = QFileDialog.getOpenFileName(self, "Cargar Prompts", "", "JSON (*.json)")
        if file:
            with open(file, 'r') as f:
                self.prompts = json.load(f)
            self.prompt_editor.setPlainText(json.dumps(self.prompts, indent=4))

    def start_review(self):
        if not self.file_path:
            QMessageBox.warning(self, "Error", "Selecciona un archivo primero.")
            return
        try:
            self.prompts = json.loads(self.prompt_editor.toPlainText())
        except:
            QMessageBox.warning(self, "Error", "Prompts no válidos (debe ser JSON).")
            return

        self.worker = IAWorker(self.file_path, int(self.num_keys_edit.text()), int(self.num_articles_edit.text()),
                               self.model_combo.currentText(), self.prompts, self.manual_checkbox.isChecked(),
                               self.output_combo.currentText())
        self.worker.progress.connect(self.progress_bar.setValue)
        self.worker.result.connect(lambda res: QMessageBox.info(self, "Éxito", "Revisión completada."))
        self.worker.error.connect(lambda err: QMessageBox.warning(self, "Error", err))
        self.worker.start()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())